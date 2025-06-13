import sqlite3
from database import get_db_connection, create_database

DATABASE_FILE = "music_catalog.db"

def populate_database_from_txt(artists_file, genres_file, songs_file, albums_file):
    """Заполняет БД данными из текстовых файлов"""
    conn = get_db_connection()
    if conn:
        cursor = conn.cursor()
        try:
            # Заполнение таблицы artists
            with open(artists_file, 'r', encoding='utf-8') as f:
                for line in f:
                    artist_id, artist_name, biography = line.strip().split('`')
                    cursor.execute(
                        "INSERT OR IGNORE INTO artists (id, name, biography) VALUES (?, ?, ?)",
                        (artist_id, artist_name, biography)
                    )

            # Заполнение таблицы genres
            with open(genres_file, 'r', encoding='utf-8') as f:
                for line in f:
                    genre_id, genre_name, description = line.strip().split('`')
                    cursor.execute(
                        "INSERT OR IGNORE INTO genres (id, name, description) VALUES (?, ?, ?)",
                        (genre_id, genre_name, description)
                    )

            # Заполнение таблицы albums
            with open(albums_file, 'r', encoding='utf-8') as f:
                for line in f:
                    album_id, album_name, artist_id, description = line.strip().split('`')
                    cursor.execute(
                        "INSERT OR IGNORE INTO albums (id, name, artist_id, description) VALUES (?, ?, ?, ?)",
                        (album_id, album_name, artist_id, description)
                    )

            # Заполнение таблицы songs
            with open(songs_file, 'r', encoding='utf-8') as f:
                for line in f:
                    song_id, title, artist_id, genre_id, album_id, year = line.strip().split('`')
                    cursor.execute(
                        "INSERT OR IGNORE INTO songs (id, title, artist_id, genre_id, album_id, year) VALUES (?, ?, ?, ?, ?, ?)",
                        (song_id, title, artist_id, genre_id, album_id, year)
                    )

            conn.commit()
            print("База данных успешно заполнена из TXT-файлов.")
        except sqlite3.Error as e:
            print(f"Ошибка при заполнении базы данных: {e}")
            conn.rollback()
        finally:
            conn.close()
    else:
        print("Не удалось подключиться к базе данных.")

def add_song(artist_name, title, genre_name, album_name, year):
    """Добавляет новую песню в БД"""
    conn = get_db_connection()
    if conn:
        cursor = conn.cursor()
        try:
            # Проверяем/добавляем артиста
            cursor.execute("SELECT id FROM artists WHERE name = ?", (artist_name,))
            artist = cursor.fetchone()
            if artist is None:
                cursor.execute("INSERT INTO artists (name) VALUES (?)", (artist_name,))
                artist_id = cursor.lastrowid
            else:
                artist_id = artist['id']

            # Проверяем/добавляем жанр
            cursor.execute("SELECT id FROM genres WHERE name = ?", (genre_name,))
            genre = cursor.fetchone()
            if genre is None:
                cursor.execute("INSERT INTO genres (name) VALUES (?)", (genre_name,))
                genre_id = cursor.lastrowid
            else:
                genre_id = genre['id']

            # Проверяем/добавляем альбом
            cursor.execute("SELECT id FROM albums WHERE name = ? AND artist_id = ?",
                          (album_name, artist_id))
            album = cursor.fetchone()
            if album is None:
                cursor.execute("INSERT INTO albums (name, artist_id) VALUES (?, ?)",
                             (album_name, artist_id))
                album_id = cursor.lastrowid
            else:
                album_id = album['id']

            # Добавляем песню
            cursor.execute(
                "INSERT INTO songs (title, artist_id, genre_id, album_id, year) VALUES (?, ?, ?, ?, ?)",
                (title, artist_id, genre_id, album_id, year)
            )
            conn.commit()
            print(f"Песня '{title}' успешно добавлена.")
        except sqlite3.Error as e:
            print(f"Ошибка при добавлении песни: {e}")
            conn.rollback()
        finally:
            conn.close()
    else:
        print("Не удалось подключиться к базе данных.")

def update_song(song_title, new_title=None, new_artist_name=None,
               new_genre_name=None, new_album_name=None, new_year=None):
    """Обновляет информацию о песне"""
    conn = get_db_connection()
    if not conn:
        print("Не удалось подключиться к базе данных.")
        return False

    try:
        cursor = conn.cursor()
        cursor.execute("SELECT id, artist_id FROM songs WHERE title = ?", (song_title,))
        song = cursor.fetchone()
        if not song:
            print(f"Песня с названием '{song_title}' не найдена.")
            return False

        song_id = song['id']
        current_artist_id = song['artist_id']

        if new_title:
            cursor.execute("UPDATE songs SET title = ? WHERE id = ?", (new_title, song_id))

        if new_artist_name:
            cursor.execute("SELECT id FROM artists WHERE name = ?", (new_artist_name,))
            artist = cursor.fetchone()
            if not artist:
                cursor.execute("INSERT INTO artists (name) VALUES (?)", (new_artist_name,))
                new_artist_id = cursor.lastrowid
            else:
                new_artist_id = artist['id']
            cursor.execute("UPDATE songs SET artist_id = ? WHERE id = ?", (new_artist_id, song_id))
            current_artist_id = new_artist_id

        if new_genre_name is not None:
            if new_genre_name == "":
                cursor.execute("UPDATE songs SET genre_id = NULL WHERE id = ?", (song_id,))
            else:
                cursor.execute("SELECT id FROM genres WHERE name = ?", (new_genre_name,))
                genre = cursor.fetchone()
                if not genre:
                    cursor.execute("INSERT INTO genres (name) VALUES (?)", (new_genre_name,))
                    new_genre_id = cursor.lastrowid
                else:
                    new_genre_id = genre['id']
                cursor.execute("UPDATE songs SET genre_id = ? WHERE id = ?", (new_genre_id, song_id))

        if new_album_name is not None:
            if new_album_name == "":
                cursor.execute("UPDATE songs SET album_id = NULL WHERE id = ?", (song_id,))
            else:
                cursor.execute(
                    "SELECT id FROM albums WHERE name = ? AND artist_id = ?",
                    (new_album_name, current_artist_id)
                )
                album = cursor.fetchone()
                if not album:
                    cursor.execute(
                        "INSERT INTO albums (name, artist_id) VALUES (?, ?)",
                        (new_album_name, current_artist_id)
                    )
                    new_album_id = cursor.lastrowid
                else:
                    new_album_id = album['id']
                cursor.execute("UPDATE songs SET album_id = ? WHERE id = ?", (new_album_id, song_id))

        if new_year:
            cursor.execute("UPDATE songs SET year = ? WHERE id = ?", (new_year, song_id))

        conn.commit()
        print(f"Песня '{song_title}' успешно обновлена.")
        return True
    except sqlite3.Error as e:
        print(f"Ошибка при обновлении песни: {e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def delete_song(song_title):
    conn = get_db_connection()
    if conn:
        cursor = conn.cursor()
        try:
            cursor.execute("SELECT id FROM songs WHERE title = ?", (song_title,))
            song = cursor.fetchone()
            if not song:
                return False

            cursor.execute("DELETE FROM songs WHERE id = ?", (song['id'],))
            conn.commit()
            return True
        except sqlite3.Error as e:
            conn.rollback()
            print(f"Ошибка при удалении песни: {e}")
            return False
        finally:
            conn.close()
    else:
        return False


def delete_artist(artist_name):
    """Удаляет исполнителя и все связанные данные"""
    conn = get_db_connection()
    if conn:
        cursor = conn.cursor()
        try:
            cursor.execute("SELECT id FROM artists WHERE name = ?", (artist_name,))
            artist = cursor.fetchone()
            if not artist:
                print(f"Исполнитель с именем '{artist_name}' не найден.")
                return  False

            artist_id = artist['id']
            cursor.execute("DELETE FROM songs WHERE artist_id = ?", (artist_id,))
            cursor.execute("DELETE FROM albums WHERE artist_id = ?", (artist_id,))
            cursor.execute("DELETE FROM artists WHERE id = ?", (artist_id,))
            conn.commit()
            print(f"Исполнитель '{artist_name}' и все его песни и альбомы успешно удалены.")
            return True
        except sqlite3.Error as e:
            conn.rollback()
            print(f"Ошибка при удалении исполнителя: {e}")
            return False
        finally:
            conn.close()
    else:
        print("Не удалось подключиться к базе данных.")
        return False

def delete_album(album_name):
    """Удаляет альбом из БД"""
    conn = get_db_connection()
    if conn:
        cursor = conn.cursor()
        try:
            cursor.execute("SELECT id FROM albums WHERE name = ?", (album_name,))
            album = cursor.fetchone()
            if not album:
                print(f"Альбом с названием '{album_name}' не найден.")
                return False

            album_id = album['id']
            cursor.execute("DELETE FROM songs WHERE album_id = ?", (album_id,))
            cursor.execute("DELETE FROM albums WHERE id = ?", (album_id,))
            conn.commit()
            print(f"Альбом '{album_name}' успешно удален. Ссылки на него в песнях обнулены.")
            return True
        except sqlite3.Error as e:
            conn.rollback()
            print(f"Ошибка при удалении альбома: {e}")
            return False
        finally:
            conn.close()
    else:
        print("Не удалось подключиться к базе данных.")
        return False

def search_tracks(query):
    """Поиск треков по названию, альбому или исполнителю"""
    conn = get_db_connection()
    if conn is None:
        print("Ошибка подключения к БД")
        return None
    try:
        cursor = conn.cursor()
        search_term = f'%{query}%'
        cursor.execute('''
            SELECT 
                s.id, s.title as name, 
                a.id as album_id, a.name as album_name,
                ar.id as artist_id, ar.name as artist_name
            FROM songs s
            LEFT JOIN albums a ON s.album_id = a.id
            LEFT JOIN artists ar ON s.artist_id = ar.id
            WHERE s.title LIKE ? OR a.name LIKE ? OR ar.name LIKE ?
        ''', (search_term, search_term, search_term))

        return [dict(row) for row in cursor.fetchall()]
    finally:
        conn.close()

def get_album_details(album_name):
    """Получает детальную информацию об альбоме"""
    conn = get_db_connection()
    try:
        cursor = conn.cursor()
        cursor.execute('''
            SELECT 
                a.id, a.name, a.description,
                ar.id as artist_id, ar.name as artist_name
            FROM albums a
            JOIN artists ar ON a.artist_id = ar.id
            WHERE a.name = ?
        ''', (album_name,))

        album = cursor.fetchone()
        if not album:
            return None

        cursor.execute('''
            SELECT id, title, year 
            FROM songs 
            WHERE album_id = ?
        ''', (album['id'],))

        songs = [dict(row) for row in cursor.fetchall()]
        return {
            'album': dict(album),
            'songs': songs
        }
    finally:
        conn.close()

def get_artist_albums(artist_name):
    """Получает все альбомы исполнителя"""
    conn = get_db_connection()
    try:
        cursor = conn.cursor()
        cursor.execute('''
            SELECT a.id, a.name, a.description
            FROM albums a
            JOIN artists ar ON a.artist_id = ar.id
            WHERE ar.name = ?
        ''', (artist_name,))

        return [dict(row) for row in cursor.fetchall()]
    finally:
        conn.close()

def export_songs_to_docx(filename="songs_export.docx"):
    conn = get_db_connection()
    if not conn:
        print("Не удалось подключиться к базе данных.")
        return False

    try:
        from docx import Document
        from docx.shared import Pt

        cursor = conn.cursor()
        cursor.execute('''
            SELECT 
                s.title, s.year,
                ar.name as artist,
                al.name as album,
                g.name as genre
            FROM songs s
            LEFT JOIN artists ar ON s.artist_id = ar.id
            LEFT JOIN albums al ON s.album_id = al.id
            LEFT JOIN genres g ON s.genre_id = g.id
            ORDER BY ar.name, s.year
        ''')

        songs = cursor.fetchall()
        if not songs:
            print("Нет песен для экспорта.")
            return False

        # Создаем документ
        doc = Document()
        doc.add_heading('Каталог песен', level=1)

        # Настройка стиля таблицы
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название'
        hdr_cells[1].text = 'Исполнитель'
        hdr_cells[2].text = 'Год'
        hdr_cells[3].text = 'Альбом'
        hdr_cells[4].text = 'Жанр'

        # Заполняем таблицу данными
        for song in songs:
            row_cells = table.add_row().cells
            row_cells[0].text = song['title']
            row_cells[1].text = song['artist']
            row_cells[2].text = str(song['year'])
            row_cells[3].text = song['album'] or "—"
            row_cells[4].text = song['genre'] or "—"

        # Сохраняем файл
        doc.save(filename)
        print(f"Документ сохранен как {filename}")
        return True

    except Exception as e:
        print(f"Ошибка: {e}")
        return False
    finally:
        conn.close()

def clear_database():
    try:
        conn = get_db_connection()
        if conn is None:
            print("Не удалось получить соединение с базой данных")
            return False

        cursor = conn.cursor()

        # Отключаем foreign keys
        cursor.execute("PRAGMA foreign_keys = OFF;")

        # Чистим таблицы
        cursor.execute("DELETE FROM songs;")
        cursor.execute("DELETE FROM albums;")
        cursor.execute("DELETE FROM artists;")
        cursor.execute("DELETE FROM genres;")

        conn.commit()

        # Включаем обратно
        cursor.execute("PRAGMA foreign_keys = ON;")
        print("База данных успешно очищена")
        return True

    except Exception as e:
        print(f"Ошибка при очистке базы данных: {e}")
        return False

    finally:
        if conn:
            conn.close()
# Инициализация БД
create_database(DATABASE_FILE)


# Заполняем базу данных из TXT-файлов
# artists_file = "artists.txt"
# genres_file = "genres.txt"
# songs_file = "songs.txt"
# albums_file = "albums.txt"

# populate_database_from_txt(artists_file, genres_file, songs_file, albums_file)

# # Добавляем песни в базу данных
# # add_song("МакSим", "Знаешь ли ты", "Поп-музыка" , "Трудный возраст",2007)
# # add_song("ВИА 'Поющие сердца' ","Кто тебе сказал", "Русская эстрада", "Гигант", 1975)
# add_song("1234", "321", "zx1", "asd1", 123)
# # delete_song("ytrew1")
# # delete_song("qwerty")
# # delete_artist("123")
# update_song("321", "b", "q", "j", )
# delete_album("Альбом2")
search_tracks("")
